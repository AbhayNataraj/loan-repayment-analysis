from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "UCUCI_Loan_Repayment_Analysis_Report.docx"
CROP = ROOT / "tmp" / "pdfs" / "report_v4_crops"
NOTEBOOK = ROOT / "tmp" / "pdfs" / "notebook_figures"

FONT = "Arial"
INK = RGBColor(31, 41, 51)
NAVY = RGBColor(16, 42, 67)
MUTED = RGBColor(91, 103, 112)
FILL = "E7F5EE"
RULE = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = RULE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            set_cell_border(row.cells[idx])


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Title", 22, NAVY, 0, 6),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 17, NAVY, 14, 6),
        ("Heading 2", 14, NAVY, 10, 4),
        ("Heading 3", 12, NAVY, 8, 3),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if name != "Subtitle":
            style.font.bold = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.add_run("UCUCI Bank Loan Repayment Analysis")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Improving repayment performance and reducing default exposure")
    p = doc.add_paragraph()
    r = p.add_run("Format validation note: ")
    r.bold = True
    p.add_run(
        "editable report text uses Arial at 12 pt or larger. Chart labels are embedded in source images and may render smaller because they are part of the visuals."
    )


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    header = table.rows[0].cells
    for cell, value in zip(header, headers):
        set_cell_shading(cell, FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)
            set_cell_margins(cell)
            set_cell_border(cell)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.4) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED


def add_finding(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"Insight - {label}: ")
    r.bold = True
    p.add_run(text)


def page_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("Business Context", level=1)
    add_para(
        doc,
        "UCUCI Bank has a large retail loan portfolio covering debt consolidation, credit card, home improvement, small business, personal, and other loan purposes. While the portfolio-level repayment rate looks healthy, defaults and delinquency are not evenly distributed.",
    )
    add_para(
        doc,
        "The business problem is to identify which loan characteristics are linked with weak repayment, then use those patterns to improve future lending decisions and prioritise recovery for active loans before they become charged off.",
        True,
    )
    doc.add_heading("Business Objective", level=2)
    add_bullets(
        doc,
        [
            "Find the main drivers of completed-loan repayment rate across grade, interest rate, DTI, term, purpose, and combined risk groups.",
            "Use completed-loan behaviour to infer which active loans are more likely to become repayment problems.",
            "Recommend practical policy and recovery actions that can improve repayment performance and reduce default exposure.",
        ],
    )

    doc.add_heading("Dataset and Key Metrics", level=1)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Dataset", "UCUCI Bank Loan Dataset"],
            ["Shape", "887,379 rows and 57 columns"],
            ["Loan record check", "No duplicate loan IDs and no duplicate member IDs were found."],
            ["Closed loans", "Fully Paid, Charged Off, Default"],
            ["Active loans", "Current, In Grace Period, Late (16-30 days), Late (31-120 days)"],
        ],
        [1.55, 4.95],
    )
    add_table(
        doc,
        ["Metric", "Meaning"],
        [
            ["Loan Repayment Rate", "Total amount repaid divided by total funded amount. LRR can be above 100% because repayments include interest and fees in addition to principal."],
            ["Default Rate", "Charged-off loans divided by total loans."],
            ["Active Delinquency Rate", "Active loans in grace or late status divided by active loans."],
            ["At-Risk Outstanding", "Outstanding principal on active delinquent loans."],
            ["Practical Recovery Rate", "Recoveries divided by funded amount for charged-off loans. This is calculated only for charged-off loans because charged-off loans have zero outstanding principal in this dataset."],
        ],
        [1.75, 4.75],
    )

    doc.add_heading("EDA Approach and Hypotheses", level=1)
    doc.add_heading("EDA steps taken", level=2)
    add_bullets(
        doc,
        [
            "Validated loan count, duplicate IDs, status groups, and the fields needed for repayment, delinquency, exposure, and recovery metrics.",
            "Created closed-loan views for grade, sub-grade, interest rate, DTI, income, term, public records, purpose, and combined grade-term risk.",
            "Created active-loan views for delinquency, outstanding principal, last-payment recency, revolving utilization, recent inquiries, purpose exposure, and operational risk tiers.",
        ],
    )
    doc.add_heading("Hypotheses used in the report", level=2)
    add_table(
        doc,
        ["Closed-loan hypotheses", "Active-loan hypotheses"],
        [
            ["Lower credit grades should show weaker LRR.", "Lower grades should show higher active delinquency."],
            ["Higher interest rates should reduce repayment performance.", "High revolving utilization should signal repayment pressure."],
            ["Low income with high DTI should weaken affordability.", "Recent credit inquiries should indicate rising credit stress."],
            ["Public records and 60-month terms should reduce repayment.", "Stale last payments should increase recovery urgency."],
            ["Purpose, sub-grade, and loan age should add risk detail.", "Recovery priority should combine delinquency, exposure, recency, utilization, and inquiries."],
        ],
        [3.25, 3.25],
    )

    page_break(doc)
    doc.add_heading("Portfolio Health Snapshot", level=1)
    add_figure(doc, CROP / "kpi_row.png", "Portfolio KPI row from updated dashboard.", 6.4)
    add_figure(doc, CROP / "portfolio_mix.png", "Portfolio mix by loan status.", 6.1)
    add_figure(doc, CROP / "at_risk_overview.png", "At-risk exposure by grade.", 6.1)
    add_finding(
        doc,
        "portfolio health",
        "The portfolio has 887,379 loans and completed-loan LRR of 100.3%. The default rate is 5.08%, recovery rate on charged-off loans is 6.1%, and at-risk outstanding principal is concentrated in grades C, D, and E.",
    )

    page_break(doc)
    doc.add_heading("Closed Loan Drivers", level=1)
    add_figure(doc, CROP / "closed_grade.png", "Repayment rate declines by lower grade.", 6.4)
    add_finding(doc, "hypothesis assumed", "Lower grades should have lower repayment rates because grade captures borrower credit quality. Result: A and B loans are above 104% LRR, while G falls to 89.5%, so grade remains the clearest historical repayment signal.")
    add_figure(doc, CROP / "closed_interest.png", "Higher interest loans have lower LRR.", 6.4)
    add_finding(doc, "hypothesis assumed", "Higher interest should indicate greater borrower risk and repayment burden. Result: loans below 10% interest have 103.7% LRR, while the 25-30% band drops to 82.0%; higher pricing alone is not enough to offset high-risk lending.")
    add_figure(doc, CROP / "closed_public_records.png", "Public records reduce repayment.", 6.4)
    add_finding(doc, "hypothesis assumed", "Borrowers with public records should repay worse because prior credit events indicate financial stress. Result: LRR declines from 100.7% for borrowers with no public record to about 94.4%-94.8% once records appear.")
    add_figure(doc, CROP / "closed_income_dti.png", "Low income plus high DTI weakens repayment.", 6.4)
    add_finding(doc, "hypothesis assumed", "High DTI should be more dangerous when income is low because less repayment capacity is available. Result: borrowers earning <=25K in the 30-40 DTI band show 86.4% LRR, while stronger-income groups generally repay better at the same DTI level.")

    page_break(doc)
    doc.add_heading("Closed Loans: Purpose, Term, and Granular EDA", level=1)
    add_figure(doc, CROP / "closed_grade_term.png", "Lower grade plus longer term is weakest.", 6.4)
    add_finding(doc, "hypothesis assumed", "Longer terms should weaken repayment when combined with lower credit grades because uncertainty and repayment burden last longer. Result: G-grade 60-month loans show 88.9% LRR, while A-grade 36-month loans show 104.3%.")
    add_figure(doc, NOTEBOOK / "cell_077_0.png", "Python EDA: sub-grade charge-off pattern.", 6.1)
    add_finding(doc, "hypothesis assumed", "Sub-grade should reveal risk hidden inside broad grades. Result: Python EDA shows charge-off rate generally increases as sub-grade quality declines, so sub-grade refines monitoring beyond broad A to G grades.")
    add_figure(doc, NOTEBOOK / "cell_082_0.png", "Python EDA: loan age and charge-off risk.", 6.1)
    add_finding(doc, "hypothesis assumed", "Charge-off risk should appear before the loan reaches very late stages. Result: Python EDA shows charge-off risk is highest in early loan-age buckets, especially around 6-12 months.")

    page_break(doc)
    doc.add_heading("Closed Loan Driver Summary", level=1)
    add_figure(doc, CROP / "closed_dashboard.png", "Updated closed loan repayment driver dashboard.", 6.4)
    add_finding(
        doc,
        "closed loans",
        "Weak repayment is concentrated by lower grade, high interest, public records, lower income plus high DTI, and weaker long-term grade combinations. These patterns become the basis for active-loan monitoring.",
    )
    doc.add_heading("Applying Closed-Loan Learning to Active Loans", level=2)
    add_table(
        doc,
            ["Closed-loan pattern", "Active-loan use"],
            [
                ["Lower grades had weaker LRR.", "Use grade as the first early-warning filter."],
            ["High interest bands had weaker LRR.", "Monitor high-rate active loans earlier and review pricing-risk overlap."],
            ["High DTI, public records, and 60-month terms added risk.", "Flag affordability and credit-history stress before accounts become severely late."],
            ["Purpose changed repayment behaviour.", "Combine purpose risk with outstanding principal before assigning recovery priority."],
        ],
        [2.65, 3.85],
    )

    page_break(doc)
    doc.add_heading("Active Loan Risk and Recovery Priority", level=1)
    add_figure(doc, CROP / "active_grade.png", "Active delinquency rises sharply in lower grades.", 6.3)
    add_finding(doc, "hypothesis assumed", "Active delinquency should rise in lower grades if historical repayment risk is already appearing in the live portfolio. Result: delinquency rises from 1.0% in grade A to 11.6% in grade G.")
    add_figure(doc, CROP / "active_utilization.png", "High revolving utilization raises active risk.", 6.3)
    add_finding(doc, "hypothesis assumed", "High revolving utilization should signal repayment pressure because borrowers have less unused credit capacity. Result: active delinquency increases from 2.2% in the 0-30% utilization band to 5.4% in the 100%+ band.")
    add_figure(doc, CROP / "active_recency.png", "Stale last payments increase recovery priority.", 6.3)
    add_finding(doc, "hypothesis assumed", "Stale or weak payment activity should raise recovery priority when the loan is already delinquent. Result: at-risk outstanding is largest in 0-1 month at 12.35 Cr, followed by 2-3 months at 7.90 Cr and no-payment-date loans at 3.31 Cr.")
    add_figure(doc, CROP / "active_inquiries.png", "Recent credit inquiries weaken active repayment.", 6.3)
    add_finding(doc, "hypothesis assumed", "Recent credit inquiries should indicate rising credit demand and stress. Result: delinquency increases from 2.7% for borrowers with no recent inquiries to 6.8% for borrowers with 5+ inquiries.")
    add_figure(doc, CROP / "active_scatter.png", "Exposure and delinquency prioritize recovery.", 6.3)

    page_break(doc)
    doc.add_heading("Active Risk and Recovery Summary", level=1)
    add_figure(doc, CROP / "active_dashboard.png", "Updated active loan risk and recovery prioritization dashboard.", 6.4)
    add_finding(
        doc,
        "active portfolio",
        "The active-risk view links delinquency with exposure. Recovery should start where risk and outstanding principal are both high, rather than chasing only the highest delinquency percentage.",
    )
    doc.add_heading("Priority Order Suggested by the Analysis", level=2)
    add_bullets(
        doc,
        [
            "First priority: delinquent C, D, and E grade loans with high outstanding principal.",
            "Second priority: loans in the 0-1 month and 2-3 month recency bands before exposure becomes harder to recover.",
            "Third priority: active loans with high revolving utilization, 3+ recent inquiries, or high-risk purpose segments, reviewed after exposure size is considered.",
        ],
    )

    doc.add_heading("Recommendations", level=1)
    add_table(
        doc,
        ["Recommendation", "Business action"],
        [
            ["Prioritise active recovery", "Start with delinquent C, D, and E loans, then rank by outstanding principal. These grades hold about 16.7 Cr of at-risk principal combined."],
            ["Add payment-recency monitoring", "Use last payment recency to trigger faster recovery action. The 0-1 month and 2-3 month bands hold 12.35 Cr and 7.90 Cr of at-risk exposure."],
            ["Tighten future lending rules", "Add stricter review when lower grades, high DTI, high interest rates, public records, and 60-month terms overlap. G-grade 60-month loans show 88.9% LRR and the 25-30% interest band shows 82.0% LRR."],
            ["Strengthen affordability checks", "Assess income with DTI, revolving utilization, and recent inquiries. Active delinquency reaches 5.4% at 100%+ utilization and 6.8% for borrowers with 5+ recent inquiries."],
            ["Act earlier in the loan lifecycle", "Because charge-off risk appears early, intervention should start before loans become severely late."],
        ],
        [2.05, 4.45],
    )
    add_para(
        doc,
        "Final conclusion: UCUCI should use closed-loan repayment patterns to improve approval discipline and use active-loan delinquency plus outstanding principal to guide recovery effort. The strongest opportunity is targeted action on combinations where weak repayment and high exposure meet.",
        True,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
